
import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def export_paper_to_docx():
    # Paths
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # MLAssignment root
    md_path = os.path.join(base_dir, "report", "paper.md")
    docx_path = os.path.join(base_dir, "report", "paper.docx")
    figures_dir = os.path.join(base_dir, "figures")

    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # --- Style Configuration for "Normal" (Body Text) ---
    # 宋体 (SimSun), 小四 (12pt), 单倍行距
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    
    # Python-docx handling for East Asian fonts
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Iterate and Parse
    code_block_mode = False
    code_content = []

    for line in lines:
        line_stripped = line.strip()
        
        # 1. Handle Code Blocks
        if line_stripped.startswith("```"):
            if code_block_mode:
                # End of code block
                p = doc.add_paragraph("\n".join(code_content))
                p.style = 'No Spacing' # Use a compact style for code
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(10)
                code_content = []
                code_block_mode = False
            else:
                # Start of code block
                code_block_mode = True
            continue
        
        if code_block_mode:
            code_content.append(line.rstrip())
            continue

        # 2. Handle Images: ![Alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line_stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            # Fix path: replace "figures/" with actual absolute path
            img_name = os.path.basename(img_rel_path)
            img_full_path = os.path.join(figures_dir, img_name)
            
            if os.path.exists(img_full_path):
                try:
                    doc.add_picture(img_full_path, width=Inches(5.5))
                    # Add caption centered
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption.runs[0].font.italic = True
                    caption.runs[0].font.size = Pt(10)
                except Exception as e:
                    print(f"Error adding image {img_full_path}: {e}")
            else:
                doc.add_paragraph(f"[Image not found: {img_name}]")
            continue

        # 3. Handle Headers
        if line_stripped.startswith('#'):
            level = len(line_stripped.split()[0])
            text = line_stripped.lstrip('#').strip()
            # Title (Level 1) or Chapters (Level 2+)
            if level == 1:
                p = doc.add_heading(text, 0) # Title
            else:
                p = doc.add_heading(text, level=level-1)
            continue
        
        # 4. Handle Lists (Simple)
        if line_stripped.startswith('* ') or line_stripped.startswith('- '):
            text = line_stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            # Apply Chinese font manually to list items as they might not inherit 'Normal' eastAsia setting automatically in some viewers
            for run in p.runs:
                 run.font.name = 'SimSun'
                 run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                 run.font.size = Pt(12)
            continue
        
        # 5. Handle Ordered Lists (Simple numeric)
        if re.match(r'^\d+\.', line_stripped):
            text = list(line_stripped.partition('. '))[2].strip()
            p = doc.add_paragraph(text, style='List Number')
             # Apply Chinese font manually
            for run in p.runs:
                 run.font.name = 'SimSun'
                 run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                 run.font.size = Pt(12)
            continue

        # 6. Normal Paragraphs (skip empty lines)
        if not line_stripped:
            continue
            
        p = doc.add_paragraph(line_stripped)
        # Ensure the font is applied
        for run in p.runs:
             run.font.name = 'SimSun'
             run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
             run.font.size = Pt(12)

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == "__main__":
    export_paper_to_docx()
